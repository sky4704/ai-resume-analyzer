from docx import Document
from docx.shared import Inches
from datetime import datetime
import os
import uuid
import json
from docx.oxml import OxmlElement
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_TAB_ALIGNMENT
from docx.oxml.ns import qn
from gemini_ai import client
import tempfile


def create_resume_doc(data):

    doc = Document()

    style = doc.styles["Normal"]

    font = style.font

    font.name = "Calibri"

    font.size = Pt(11)

    section = doc.sections[0]

    section.top_margin = Inches(0.6)

    section.bottom_margin = Inches(0.6)

    section.left_margin = Inches(0.7)

    section.right_margin = Inches(0.7)

    # =========================
    # HEADER
    # =========================

    name = data.get("name", "")

    email = data.get("email", "")

    phone = data.get("phone", "")

    address = data.get("address", "")

    linkedin = data.get("linkedin", "")

    github = data.get("github", "")

    target_role = data.get("target_role", "")

    summary = data.get("summary", "")

    skills = data.get("skills", [])

    experience = data.get("experience", [])

    projects = data.get("projects", [])

    certifications = data.get("certifications",[])

    education = data.get("education", "")

    declaration = data.get("declaration","")

    # =========================
    # HEADER
    # =========================

    heading = doc.add_paragraph()

    heading.alignment = (
        WD_PARAGRAPH_ALIGNMENT.LEFT
    )

    heading.paragraph_format.space_after = Pt(0)

    name_run = heading.add_run(name.title())

    name_run.bold = True

    name_run.font.size = Pt(24)

    name_run.font.name = "Calibri"

    # =========================
    # ROLE TITLE
    # =========================

    role_para = doc.add_paragraph()

    role_para.alignment = (WD_PARAGRAPH_ALIGNMENT.LEFT)

    role_para.paragraph_format.space_after = Pt(5)

    role_run = role_para.add_run(target_role.title())

    role_run.italic = True

    role_run.font.size = Pt(13)

    role_run.font.name = "Calibri"

    # =========================
    # CONTACT INFO
    # =========================

    contact = doc.add_paragraph()

    contact.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    contact_format = contact.paragraph_format

    contact_format.space_after = Pt(21)

    contact_format.tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)

    contact.add_run(email)

    if linkedin:
        contact.add_run(
            f"\t{linkedin}"
        )

    contact.add_run("\n")
    contact.add_run(phone)
    if github:
        contact.add_run(
            f"\t{github}"
        )

    contact.add_run("\n")

    addr_run = contact.add_run(address)

    addr_run.font.size = Pt(10.5)


    # =========================
    # PROFESSIONAL SUMMARY
    # =========================

    add_section_heading(doc, "Professional Summary")

    doc.add_paragraph(summary)

    # =========================
    # SKILLS
    # =========================

    add_section_heading(doc, "Skills")

    skills_para = doc.add_paragraph()

    skills_para.add_run(
        ", ".join(skills)
    )

    # =========================
    # EXPERIENCE
    # =========================

    add_section_heading(doc, "Experience")

    for exp in experience:

        # Support old string format safely

        if isinstance(exp, str):

            doc.add_paragraph(
                exp,
                style="List Bullet"
            )

            continue

        title = exp.get("title", "")

        company = exp.get("company", "")

        dates = exp.get("dates", "")

        description = exp.get(
            "description",
            ""
        )

        p = doc.add_paragraph()

        para_format = p.paragraph_format

        para_format.space_after = Pt(4)

        run = p.add_run(
            f"{title} | {company}"
        )

        run.bold = True

        if dates:

            p.add_run(f" ({dates})")

        doc.add_paragraph(
            description,
            style="List Bullet"
        )

    # =========================
    # PROJECTS
    # =========================

    add_section_heading(doc, "Projects")

    for proj in projects:

        # Support old string format

        if isinstance(proj, str):

            doc.add_paragraph(
                proj,
                style="List Bullet"
            )

            continue

        title = proj.get("title", "")

        description = proj.get(
            "description",
            ""
        )

        p = doc.add_paragraph()

        para_format = p.paragraph_format

        para_format.space_after = Pt(4)

        run = p.add_run(title)

        run.bold = True

        doc.add_paragraph(
            description,
            style="List Bullet"
    )
        
    # =========================
    # CERTIFICATES
    # =========================

    add_section_heading(doc,"Certifications")
    
    for cert in certifications:

        doc.add_paragraph(
            cert,
            style="List Bullet"
        )

    # =========================
    # EDUCATION
    # =========================

    add_section_heading(doc, "Education")

    doc.add_paragraph(education)

    # =========================
    # Declaration
    # =========================

    if declaration:

        add_section_heading(
            doc,
            "Declaration"
        )

        para = doc.add_paragraph(
            declaration
        )

        para.style = doc.styles["Normal"]

    # =========================
    # SAVE FILE
    # =========================

    os.makedirs(
        "generated_resumes",
        exist_ok=True
    )

    filename = f"{name.replace(" ", "").lower()}.docx"

    temp_file = tempfile.NamedTemporaryFile(
        delete=False,
        suffix=".docx"
    )

    filepath = temp_file.name

    doc.save(filepath)

    return filepath


def add_horizontal_line(paragraph):

    p = paragraph._p

    pPr = p.get_or_add_pPr()

    border = OxmlElement("w:pBdr")

    bottom = OxmlElement("w:bottom")

    bottom.set(qn("w:val"), "single")

    bottom.set(qn("w:sz"), "3")

    bottom.set(qn("w:space"), "4")

    bottom.set(qn("w:color"), "A0A0A0")

    border.append(bottom)

    pPr.append(border)



def add_section_heading(doc, title):

    para = doc.add_paragraph()

    para.space_before = Pt(14)

    para.space_after = Pt(4)

    run = para.add_run(title.upper())

    run.bold = True

    run.font.size = Pt(12)

    run.font.name = "Calibri"

    add_horizontal_line(para)

    return para




def enhance_resume_content(data):

    prompt = f"""
        You are an elite, results-oriented professional resume writer.
        Your task is to rewrite the provided resume data to be highly impactful, punchy, and clean.

        CRITICAL WRITING RULES:
        - **Style**: Clear, active, and concise. Completely avoid corporate fluff and jargon (e.g., "Responsible for", "Synergized", "Utilized cutting-edge solutions"). 
        - **Metrics**: Prioritize business impact and metrics where possible, but keep skills 100% realistic.
        - **Length Constraints**:
          - Summary: Maximum 3 sentences. Focus on core expertise and value.
          - Experience & Project Descriptions: Must be formatted as exactly 3 distinct, high-impact bullet points per item. Each bullet point must be under 20 words.

        Return ONLY a raw, valid JSON object matching the exact structure below. Do not include markdown blocks, wrappers, or extra text.

        FORMAT:

            {{
                "summary": "",

                "experience": [
                    {{
                        "title": "",
                        "company": "",
                        "dates": "",
                        "description": ""
                    }}
                ],

                "projects": [
                    {{
                        "title": "",
                        "description": ""
                    }}
                ]
            }}

        DATA:
        {json.dumps(data)}
        """

    try:

        response = client.models.generate_content(
            model="gemini-3-flash-preview",
            contents=prompt
        )

        content = response.text.strip()

        start = content.find("{")

        end = content.rfind("}") + 1

        cleaned = content[start:end]

        return json.loads(cleaned)

    except Exception as e:

        print("RESUME BUILDER AI ERROR:", e)

        return {
            "summary": data.get(
                "summary",
                ""
            ),
            "experience": data.get(
                "experience",
                []
            ),
            "projects": data.get(
                "projects",
                []
            )
        }